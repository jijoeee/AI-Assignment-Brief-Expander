import customtkinter as ctk
import tkinter as tk
from tkinter import filedialog, messagebox
import threading
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import json
import os
import PyPDF2
import webbrowser
import re

CONFIG_FILE = "app_config.json"

class AssignmentExpanderApp(ctk.CTk):
    def __init__(self):
        super().__init__()

        # Window Title
        self.title("AI Assignment Brief Expander V1.0 (Powered by Gemini)")
        self.geometry("900x820")
        ctk.set_appearance_mode("dark")
        ctk.set_default_color_theme("blue")

        self.api_key = self.load_api_key()
        self.pdf_text = ""
        self.mermaid_code = ""
        self.clean_text = ""

        # --- UI Layout ---
        self.grid_columnconfigure(0, weight=1)
        self.grid_columnconfigure(1, weight=1)
        self.grid_columnconfigure(2, weight=1)
        
        # API Setting
        self.btn_settings = ctk.CTkButton(self, text="⚙️ API Setting", command=self.open_settings, width=120, fg_color="gray")
        self.btn_settings.grid(row=0, column=2, padx=20, pady=(10, 0), sticky="e")

        # Input Section
        self.lbl_brief = ctk.CTkLabel(self, text="Paste Assignment Brief Here:", font=("Helvetica", 14, "bold"))
        self.lbl_brief.grid(row=1, column=0, columnspan=3, padx=20, pady=(10, 5), sticky="w")
        
        self.txt_brief = ctk.CTkTextbox(self, height=100)
        self.txt_brief.grid(row=2, column=0, columnspan=3, padx=20, pady=(0, 10), sticky="ew")

        # --- PDF Context Injector ---
        self.pdf_frame = ctk.CTkFrame(self, fg_color="transparent")
        self.pdf_frame.grid(row=3, column=0, columnspan=3, padx=20, pady=5, sticky="ew")
        
        self.btn_upload_pdf = ctk.CTkButton(self.pdf_frame, text="📄 Attach PDF (Rubric/Syllabus)", command=self.upload_pdf, fg_color="#8e44ad", hover_color="#9b59b6")
        self.btn_upload_pdf.pack(side="left", padx=(0, 10))
        
        self.lbl_pdf_name = ctk.CTkLabel(self.pdf_frame, text="No PDF attached", text_color="gray")
        self.lbl_pdf_name.pack(side="left", padx=10)

        self.btn_remove_pdf = ctk.CTkButton(self.pdf_frame, text="❌ Remove", command=self.remove_pdf, state="disabled", width=80, fg_color="#c0392b", hover_color="#e74c3c")
        self.btn_remove_pdf.pack(side="right")

        # --- Options Section (With Headers) ---
        self.lbl_type = ctk.CTkLabel(self, text="Assignment Type:", font=("Helvetica", 12, "bold"))
        self.lbl_type.grid(row=4, column=0, padx=20, pady=(10, 0), sticky="w")

        self.lbl_words = ctk.CTkLabel(self, text="Target Words (Max 5000):", font=("Helvetica", 12, "bold"))
        self.lbl_words.grid(row=4, column=1, padx=20, pady=(10, 0), sticky="w")

        self.lbl_format = ctk.CTkLabel(self, text="Format Template:", font=("Helvetica", 12, "bold"))
        self.lbl_format.grid(row=4, column=2, padx=20, pady=(10, 0), sticky="w")

        # Dropdowns
        self.opt_type = ctk.CTkOptionMenu(self, values=["Essay", "Report", "Case Study", "Presentation"])
        self.opt_type.grid(row=5, column=0, padx=20, pady=(5, 10), sticky="ew")

        # UPGRADED TO COMBOBOX: Includes 100, 250, and allows manual typing
        self.cmb_words = ctk.CTkComboBox(self, values=["100", "250", "500", "1000", "1500", "2500", "3000", "5000"])
        self.cmb_words.set("1500") # Default value
        self.cmb_words.grid(row=5, column=1, padx=20, pady=(5, 10), sticky="ew")

        self.opt_format = ctk.CTkOptionMenu(self, values=["Standard Format", "APA 7th Edition", "MLA 9th Edition"])
        self.opt_format.grid(row=5, column=2, padx=20, pady=(5, 10), sticky="ew")

        # --- Action Buttons ---
        self.btn_generate = ctk.CTkButton(self, text="🚀 Generate Roadmap & Mind-Map", command=self.start_generation_thread, height=45, font=("Helvetica", 14, "bold"))
        self.btn_generate.grid(row=6, column=0, columnspan=3, padx=20, pady=15)

        # --- Output Section ---
        self.txt_output = ctk.CTkTextbox(self, height=200, state="disabled")
        self.txt_output.grid(row=7, column=0, columnspan=3, padx=20, pady=(0, 10), sticky="nsew")
        self.grid_rowconfigure(7, weight=1) 

        # Action Bar (Export & Visuals)
        self.btn_mindmap = ctk.CTkButton(self, text="🧠 View Visual Mind-Map", command=self.open_mindmap, state="disabled", fg_color="#d35400", hover_color="#e67e22")
        self.btn_mindmap.grid(row=8, column=0, columnspan=2, padx=20, pady=(0, 20), sticky="ew")

        self.btn_export = ctk.CTkButton(self, text="💾 Export to Word Doc", command=self.export_to_docx, state="disabled", fg_color="#27ae60", hover_color="#2ecc71")
        self.btn_export.grid(row=8, column=2, padx=20, pady=(0, 20), sticky="ew")

    # --- Settings Management ---
    def load_api_key(self):
        if os.path.exists(CONFIG_FILE):
            try:
                with open(CONFIG_FILE, "r") as f:
                    return json.load(f).get("gemini_api_key", "")
            except Exception:
                return ""
        return ""

    def save_api_key(self, key):
        self.api_key = key
        with open(CONFIG_FILE, "w") as f:
            json.dump({"gemini_api_key": key}, f)

    def open_settings(self):
        settings_win = ctk.CTkToplevel(self)
        settings_win.title("API Setting")
        settings_win.geometry("450x200")
        settings_win.grab_set() 
        txt_key = ctk.CTkEntry(settings_win, width=400, show="*")
        txt_key.insert(0, self.api_key)
        txt_key.pack(pady=30, padx=20)
        ctk.CTkButton(settings_win, text="Save Key", command=lambda: [self.save_api_key(txt_key.get().strip()), settings_win.destroy()]).pack()

    # --- PDF Management ---
    def upload_pdf(self):
        file_path = filedialog.askopenfilename(filetypes=[("PDF Files", "*.pdf")])
        if file_path:
            try:
                with open(file_path, "rb") as file:
                    reader = PyPDF2.PdfReader(file)
                    text = ""
                    for page in reader.pages:
                        text += page.extract_text() + "\n"
                self.pdf_text = text[:15000] 
                file_name = os.path.basename(file_path)
                
                self.lbl_pdf_name.configure(text=f"Loaded: {file_name}", text_color="green")
                self.btn_remove_pdf.configure(state="normal") 
            except Exception as e:
                messagebox.showerror("PDF Error", f"Could not read PDF:\n{e}")

    def remove_pdf(self):
        self.pdf_text = ""
        self.lbl_pdf_name.configure(text="No PDF attached", text_color="gray")
        self.btn_remove_pdf.configure(state="disabled")

    # --- Core AI Logic ---
    def start_generation_thread(self):
        if not self.api_key:
            messagebox.showwarning("API Key Missing", "Please click 'API Setting' and enter your Gemini API key.")
            return

        brief = self.txt_brief.get("1.0", "end-1c").strip()
        if not brief: return

        # --- CUSTOM INPUT VALIDATION ---
        try:
            # Try to turn whatever they typed (or selected) into an integer
            word_count = int(self.cmb_words.get().strip())
            
            if word_count > 5000:
                messagebox.showwarning("Limit Exceeded", "Maximum word count is 5000. Please enter a smaller number.")
                return
            if word_count < 50:
                messagebox.showwarning("Invalid Input", "Please enter a realistic word count (minimum 50).")
                return
        except ValueError:
            # If they typed letters like "five hundred", it will catch the error here
            messagebox.showwarning("Invalid Input", "Please enter a valid number for the Target Words.")
            return

        self.btn_generate.configure(state="disabled", text="⏳ Analyzing Brief & Rubric...")
        self.txt_output.configure(state="normal")
        self.txt_output.delete("1.0", "end")
        self.txt_output.insert("1.0", "Generating roadmap and visual connections...")
        self.txt_output.configure(state="disabled")
        
        # Pass the validated word_count to the generator
        threading.Thread(target=self.generate_outline, args=(brief, word_count), daemon=True).start()

    def generate_outline(self, brief, word_count):
        doc_type = self.opt_type.get()
        
        context_block = f"\n\nPROFESSOR's RUBRIC/SYLLABUS DATA:\n{self.pdf_text}\n(Ensure the outline strictly adheres to the requirements in this rubric if provided.)" if self.pdf_text else ""

        prompt = f"""
        You are an expert academic writing assistant. A student has this assignment: "{brief}"
        Type: {doc_type} | Target Word Count: {word_count} words
        {context_block}
        
        TASK 1: Generate a detailed, structured text roadmap. Include Suggested Title, Introduction Angle, Section Breakdown (with suggested word counts for each section adding up to {word_count}), Key Points, Keywords, and Mistakes to Avoid.
        
        TASK 2: At the VERY END of your response, you MUST generate a Mermaid.js 'mindmap' code block that visualizes this outline.
        
        CRITICAL RULES FOR MERMAID SYNTAX:
        1. NO SPECIAL CHARACTERS: Do not use colons (:), quotes ("), parentheses (), or brackets [] in the mindmap nodes. 
        2. Keep node names very short (1-5 words max).
        3. Use standard spaces for indentation.
        
        Format it EXACTLY like this:
        ```mermaid
        mindmap
          AssignmentTitle
            Introduction
              MainAngle
            BodyParagraphs
              FirstKeyPoint
              SecondKeyPoint
            Conclusion
              Summary
        ```
        """

        try:
            genai.configure(api_key=self.api_key)
            model = genai.GenerativeModel('gemini-2.5-flash')
            response = model.generate_content(prompt)
            
            full_text = response.text
            mermaid_match = re.search(r'```mermaid(.*?)```', full_text, re.DOTALL)
            
            if mermaid_match:
                raw_mermaid = mermaid_match.group(1).strip()
                safe_mermaid = raw_mermaid.replace("(", "").replace(")", "").replace(":", "").replace('"', "").replace("'", "")
                self.mermaid_code = safe_mermaid
                self.clean_text = full_text.replace(mermaid_match.group(0), "").replace("```mermaid", "").replace("```", "").strip()
            else:
                self.mermaid_code = ""
                self.clean_text = full_text

            self.after(0, self.update_ui_with_result)

        except Exception as e:
            self.after(0, lambda: messagebox.showerror("API Error", str(e)))
            self.after(0, lambda: self.btn_generate.configure(state="normal", text="🚀 Generate Roadmap & Mind-Map"))

    def update_ui_with_result(self):
        self.txt_output.configure(state="normal")
        self.txt_output.delete("1.0", "end")
        self.txt_output.insert("1.0", self.clean_text)
        self.txt_output.configure(state="disabled")
        
        self.btn_generate.configure(state="normal", text="🚀 Generate Roadmap & Mind-Map")
        self.btn_export.configure(state="normal")
        if self.mermaid_code:
            self.btn_mindmap.configure(state="normal")

    # --- Visual Mind-Map Generator ---
    def open_mindmap(self):
        html_content = f"""
        <!DOCTYPE html>
        <html lang="en">
        <head>
            <meta charset="UTF-8">
            <title>Assignment Mind Map</title>
            <script type="module">
                import mermaid from 'https://cdn.jsdelivr.net/npm/mermaid@10/dist/mermaid.esm.min.mjs';
                mermaid.initialize({{ startOnLoad: true, theme: 'default' }});
            </script>
            <style> 
                body {{ 
                    margin: 0; 
                    padding: 0; 
                    background-color: #f4f4f9; 
                    font-family: sans-serif; 
                    display: flex; 
                    justify-content: center; 
                    align-items: center; 
                    height: 100vh; 
                    overflow: hidden;
                }} 
                .mermaid {{ 
                    width: 95vw; 
                    height: 95vh; 
                    display: flex;
                    justify-content: center;
                    align-items: center;
                }}
                .mermaid svg {{ 
                    width: 100% !important; 
                    height: 100% !important; 
                    max-width: none !important; 
                }}
            </style>
        </head>
        <body>
            <div class="mermaid">
            {self.mermaid_code}
            </div>
        </body>
        </html>
        """
        temp_file = "temp_mindmap.html"
        with open(temp_file, "w", encoding="utf-8") as f:
            f.write(html_content)
        
        webbrowser.open('file://' + os.path.realpath(temp_file))

    # --- Perfect Format Engine ---
    def export_to_docx(self):
        if not self.clean_text: return
        file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")])
        if not file_path: return

        format_style = self.opt_format.get()
        doc = Document()

        if format_style in ["APA 7th Edition", "MLA 9th Edition"]:
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(12)
            style.paragraph_format.line_spacing = 2.0
            style.paragraph_format.space_after = Pt(0)

        if format_style == "APA 7th Edition":
            doc.add_paragraph('\n\n\n\n') 
            title = doc.add_paragraph("[Assignment Title]\n[Your Name]\n[University Name]\n[Course Code & Name]\n[Professor's Name]\n[Date]", style='Normal')
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            title.runs[0].bold = True
            doc.add_page_break()
        elif format_style == "MLA 9th Edition":
            doc.add_paragraph("[Your Name]\n[Professor's Name]\n[Course Name]\n[Date]", style='Normal')
            title = doc.add_paragraph("[Assignment Title]", style='Normal')
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        else:
            doc.add_heading('AI Assignment Roadmap', 0)

        doc.add_paragraph(self.clean_text)

        try:
            doc.save(file_path)
            messagebox.showinfo("Success", f"Formatted ({format_style}) document saved!")
        except Exception as e:
            messagebox.showerror("Export Error", str(e))

if __name__ == "__main__":
    app = AssignmentExpanderApp()
    app.mainloop()
